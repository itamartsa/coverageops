"""
Word Document Report Generator – Radio Cross-Section (חתך רדיו)
"""
import io
from datetime import datetime
from typing import List

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.coverage_engine import CrossSectionStats, DeadZone


# ── Color palette ────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0F, 0x29, 0x4D)
ACCENT_BLUE = RGBColor(0x1A, 0x56, 0xDB)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRAY  = RGBColor(0xF3, 0xF4, 0xF6)

LEVEL_FILL: dict = {
    "excellent": "00B464",
    "good":      "7DFF6B",
    "medium":    "FFE600",
    "weak":      "FF8C00",
    "marginal":  "FF3B5C",
    "none":      "CCCCCC",
}
LEVEL_TEXT_DARK = {"good", "medium"}   # levels where black text is more readable

LEVEL_HE: dict = {
    "excellent": "מצוין",
    "good":      "טוב",
    "medium":    "בינוני",
    "weak":      "חלש",
    "marginal":  "שולי",
    "none":      "אין כיסוי",
}

RISK_FILL: dict = {
    "CRITICAL": "CC0000",
    "HIGH":     "FF4444",
    "MEDIUM":   "FFA500",
    "LOW":      "00AA44",
}
RISK_HE: dict = {
    "CRITICAL": "קריטי",
    "HIGH":     "גבוה",
    "MEDIUM":   "בינוני",
    "LOW":      "נמוך",
}
MODE_HE: dict = {"DTM": "טופוגרפיה (DTM)", "DSM": "תכסית + מבנים (DSM)"}


# ── XML helpers ───────────────────────────────────────────────────────────────
def _shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper().replace("#", ""))
    tc_pr.append(shd)


def _rtl_para(para) -> None:
    p_pr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.insert(0, bidi)


def _rtl_table(table) -> None:
    tbl_pr = table._tbl.get_or_add_tblPr()
    bidi   = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")
    tbl_pr.insert(0, bidi)


def _set_col_widths(table, widths_cm: List[float]) -> None:
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _cell_text(cell, text: str, bold=False, size_pt=10,
               align=WD_ALIGN_PARAGRAPH.RIGHT,
               color: RGBColor = BLACK) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    _rtl_para(para)
    run  = para.add_run(text)
    run.bold            = bold
    run.font.size       = Pt(size_pt)
    run.font.color.rgb  = color


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    _rtl_para(para)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.bold           = True
    run.font.color.rgb = NAVY if level == 1 else ACCENT_BLUE
    run.font.size      = Pt(14 if level == 1 else 12)
    # bottom border
    p_pr  = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1A56DB")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_kv_table(doc: Document, rows: List[tuple]) -> None:
    """Two-column key/value table (RTL)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    _rtl_table(table)
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        _shade_cell(row.cells[1], "EEF2FF")          # key col (right in RTL)
        _cell_text(row.cells[1], key, bold=True, size_pt=10)
        _cell_text(row.cells[0], str(val), size_pt=10)
    _set_col_widths(table, [11.0, 5.0])
    doc.add_paragraph()


def _add_stat_box(doc: Document, label: str, value: str,
                  fill_hex: str, text_white: bool = True) -> None:
    table = doc.add_table(rows=2, cols=1)
    _rtl_table(table)
    _shade_cell(table.rows[0].cells[0], fill_hex)
    _shade_cell(table.rows[1].cells[0], fill_hex)
    color = WHITE if text_white else BLACK
    _cell_text(table.rows[0].cells[0], value, bold=True, size_pt=20,
               align=WD_ALIGN_PARAGRAPH.CENTER, color=color)
    _cell_text(table.rows[1].cells[0], label, bold=False, size_pt=10,
               align=WD_ALIGN_PARAGRAPH.CENTER, color=color)
    doc.add_paragraph()


# ── Main generator ────────────────────────────────────────────────────────────
def generate_cross_section_docx(
    stats: CrossSectionStats,
    site_params: dict,
) -> bytes:
    doc = Document()

    # ── Page setup: A4, narrow margins ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin   = Cm(2)
    section.bottom_margin = Cm(2)

    # ── Header ───────────────────────────────────────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=1)
    _shade_cell(hdr_table.rows[0].cells[0], "0F294D")
    title_para = hdr_table.rows[0].cells[0].paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rtl_para(title_para)
    r1 = title_para.add_run("📡  דוח חתך רדיו\n")
    r1.bold           = True
    r1.font.size      = Pt(22)
    r1.font.color.rgb = WHITE
    r2 = title_para.add_run("ניתוח כיסוי ציר תנועה  |  CoverageOps")
    r2.font.size      = Pt(11)
    r2.font.color.rgb = RGBColor(0xA0, 0xC4, 0xFF)
    doc.add_paragraph()

    # Date + classification
    meta_para = doc.add_paragraph()
    _rtl_para(meta_para)
    meta_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    mr = meta_para.add_run(
        f"תאריך הפקה: {datetime.now().strftime('%d/%m/%Y  %H:%M')}     "
        f"|     אתר: {stats.site_name}     |     מסמך פנימי"
    )
    mr.font.size      = Pt(9)
    mr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    # ── Section 1: Analysis details ──────────────────────────────────────────
    _add_heading(doc, "1. פרטי הניתוח")
    _add_kv_table(doc, [
        ("שם אתר",              stats.site_name),
        ("מיקום אנטנה",        f"{site_params['lat']:.5f}°N,  {site_params['lon']:.5f}°E"),
        ("גובה אנטנה",         f"{site_params['ant_height']} מ' מעל קרקע"),
        ("תדר",                 f"{site_params['frequency']} MHz"),
        ("הספק שידור",         f"{site_params['tx_power']} dBm"),
        ("סף קבלה",            f"{site_params['rx_threshold']} dBm"),
        ("רדיוס מקסימלי",      f"{site_params['max_radius']} ק\"מ"),
        ("מודל חישוב",          MODE_HE.get(stats.mode, stats.mode)),
        ("אורך ציר",            f"{stats.total_length_km:.2f} ק\"מ"),
        ("נקודות ניתוח",        str(len(stats.points))),
        ("נקודות ציר (waypoints)", str(len(stats.waypoints))),
        ("זמן חישוב",           f"{stats.duration_sec}s"),
    ])

    # ── Section 2: Executive summary ─────────────────────────────────────────
    _add_heading(doc, "2. סיכום מנהלים")

    risk_fill = RISK_FILL.get(stats.risk_level, "888888")
    _add_stat_box(doc, "כיסוי כולל",          f"{stats.covered_pct}%",       "1A56DB")
    _add_stat_box(doc, "אזורי מתים",          str(len(stats.dead_zones)),    "6B46C1")
    _add_stat_box(doc, "RSSI מינימלי",        f"{stats.rssi_min} dBm",       "0E7490")
    _add_stat_box(doc, "הערכת סיכון תקשורת",
                  RISK_HE.get(stats.risk_level, stats.risk_level),
                  risk_fill, text_white=stats.risk_level != "MEDIUM")

    # Key findings
    _add_heading(doc, "ממצאים עיקריים", level=2)
    findings = [
        f"אחוז כיסוי כולל: {stats.covered_pct}% מהציר",
        f"RSSI ממוצע: {stats.rssi_avg} dBm  |  מינ': {stats.rssi_min} dBm  |  מקס': {stats.rssi_max} dBm",
        f"זוהו {len(stats.dead_zones)} אזורי מתים",
    ]
    if stats.dead_zones:
        max_dz = max(stats.dead_zones, key=lambda d: d.length_km)
        findings.append(
            f"אזור המת הארוך ביותר: {max_dz.length_km:.2f} ק\"מ "
            f"(בין {max_dz.start_km:.2f} ל-{max_dz.end_km:.2f} ק\"מ)"
        )
    for f in findings:
        p = doc.add_paragraph(style="List Bullet")
        _rtl_para(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.runs[0] if p.runs else p.add_run(f)
        if not p.runs:
            pass
        else:
            p.runs[0].text = f
            p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── Section 3: Coverage profile table ────────────────────────────────────
    _add_heading(doc, "3. פרופיל כיסוי לאורך הציר")

    # Subsample: max 150 rows in table
    pts = stats.points
    step = max(1, len(pts) // 150)
    display_pts = pts[::step]

    headers = ["סטטוס", "רמה", "RSSI (dBm)", "קו אורך", "קו רוחב", "מרחק מאנטנה (ק\"מ)", "מרחק בציר (ק\"מ)"]
    tbl = doc.add_table(rows=1 + len(display_pts), cols=7)
    tbl.style = "Table Grid"
    _rtl_table(tbl)

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        _shade_cell(hdr_row.cells[i], "0F294D")
        _cell_text(hdr_row.cells[i], h, bold=True, size_pt=9, color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, pt in enumerate(display_pts, start=1):
        row     = tbl.rows[ri]
        fill    = LEVEL_FILL.get(pt.level, "CCCCCC")
        txt_col = BLACK if pt.level in LEVEL_TEXT_DARK else (
                  BLACK if pt.level == "none" else WHITE)
        _shade_cell(row.cells[0], fill)
        _shade_cell(row.cells[1], fill)
        status_icon = "✓" if pt.level != "none" else "✗"
        data = [
            status_icon,
            LEVEL_HE.get(pt.level, pt.level),
            str(pt.rssi),
            f"{pt.lon:.5f}",
            f"{pt.lat:.5f}",
            f"{pt.dist_from_site:.2f}",
            f"{pt.dist_along:.2f}",
        ]
        for ci, val in enumerate(data):
            _cell_text(row.cells[ci], val, size_pt=8,
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       color=txt_col if ci < 2 else BLACK)

    _set_col_widths(tbl, [1.5, 2.2, 2.5, 2.8, 2.8, 3.0, 3.0])
    doc.add_paragraph()

    # ── Section 4: Dead zones ─────────────────────────────────────────────────
    _add_heading(doc, "4. ניתוח אזורי מתים")
    if not stats.dead_zones:
        p = doc.add_paragraph("לא זוהו אזורי מתים לאורך הציר.")
        _rtl_para(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x4F)
        p.runs[0].bold = True
    else:
        dz_headers = ["חומרה", "אורך (ק\"מ)", "קו אורך יציאה", "קו רוחב יציאה",
                      "קו אורך כניסה", "קו רוחב כניסה", "עד ק\"מ", "מ-ק\"מ", "#"]
        dz_tbl = doc.add_table(rows=1 + len(stats.dead_zones), cols=9)
        dz_tbl.style = "Table Grid"
        _rtl_table(dz_tbl)
        hdr_row = dz_tbl.rows[0]
        for i, h in enumerate(dz_headers):
            _shade_cell(hdr_row.cells[i], "5B0000")
            _cell_text(hdr_row.cells[i], h, bold=True, size_pt=9,
                       color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        for ri, dz in enumerate(stats.dead_zones, start=1):
            row = dz_tbl.rows[ri]
            severity = ("קריטי" if dz.length_km > 2
                        else "גבוה" if dz.length_km > 1
                        else "בינוני" if dz.length_km > 0.5 else "נמוך")
            sev_fill = ("CC0000" if dz.length_km > 2
                        else "FF4444" if dz.length_km > 1
                        else "FFA500" if dz.length_km > 0.5 else "FFD700")
            _shade_cell(row.cells[0], sev_fill)
            data = [
                severity,
                f"{dz.length_km:.3f}",
                f"{dz.end_lon:.5f}",
                f"{dz.end_lat:.5f}",
                f"{dz.start_lon:.5f}",
                f"{dz.start_lat:.5f}",
                f"{dz.end_km:.2f}",
                f"{dz.start_km:.2f}",
                str(ri),
            ]
            for ci, val in enumerate(data):
                _cell_text(row.cells[ci], val, size_pt=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           color=WHITE if ci == 0 else BLACK)
        _set_col_widths(dz_tbl, [2.0, 2.0, 2.5, 2.5, 2.5, 2.5, 1.8, 1.8, 1.0])
    doc.add_paragraph()

    # ── Section 5: Signal distribution ───────────────────────────────────────
    _add_heading(doc, "5. התפלגות עוצמת האות")
    dist_tbl = doc.add_table(rows=1 + len(stats.signal_distribution), cols=3)
    dist_tbl.style = "Table Grid"
    _rtl_table(dist_tbl)
    for i, h in enumerate(["אחוז", "נקודות", "רמה"]):
        _shade_cell(dist_tbl.rows[0].cells[i], "0F294D")
        _cell_text(dist_tbl.rows[0].cells[i], h, bold=True, size_pt=10,
                   color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    ordered_levels = ["excellent", "good", "medium", "weak", "marginal", "none"]
    for ri, lvl in enumerate(ordered_levels, start=1):
        d    = stats.signal_distribution.get(lvl, {"count": 0, "pct": 0})
        row  = dist_tbl.rows[ri]
        fill = LEVEL_FILL.get(lvl, "CCCCCC")
        _shade_cell(row.cells[2], fill)
        txt  = BLACK if lvl in LEVEL_TEXT_DARK or lvl == "none" else WHITE
        _cell_text(row.cells[2], LEVEL_HE.get(lvl, lvl), bold=True, size_pt=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color=txt)
        _cell_text(row.cells[1], str(d["count"]), size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row.cells[0], f"{d['pct']}%", size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(dist_tbl, [3.0, 3.0, 4.0])
    doc.add_paragraph()

    # ── Section 6: Operational recommendations ───────────────────────────────
    _add_heading(doc, "6. המלצות מבצעיות")
    for i, rec in enumerate(stats.recommendations, 1):
        p = doc.add_paragraph()
        _rtl_para(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(6)
        r_num = p.add_run(f"{i}. ")
        r_num.bold           = True
        r_num.font.color.rgb = ACCENT_BLUE
        r_num.font.size      = Pt(11)
        r_txt = p.add_run(rec)
        r_txt.font.size = Pt(10)
    doc.add_paragraph()

    # ── Footer line ───────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    _rtl_para(footer_para)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_tbl = doc.add_table(rows=1, cols=1)
    _shade_cell(footer_tbl.rows[0].cells[0], "0F294D")
    _cell_text(footer_tbl.rows[0].cells[0],
               f"הופק על ידי CoverageOps  |  {datetime.now().strftime('%d/%m/%Y %H:%M')}  |  מסמך פנימי",
               size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xA0, 0xC4, 0xFF))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
